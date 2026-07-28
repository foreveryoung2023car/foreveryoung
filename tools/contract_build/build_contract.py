from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "deliverables" / "软件系统开发及交付契约书_旅乘与蓝米.docx"

INK = "1F2937"
NAVY = "17365D"
MUTED = "667085"
LIGHT = "EEF2F6"
PALE = "F7F9FC"
LINE = "CBD5E1"
ACCENT = "9A6A2F"
FONT = "Hiragino Sans GB"


def set_run_font(run, size=10.5, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_page_break(doc):
    doc.add_page_break()


def add_numbering_definition(doc, fmt="decimal", text="%1."):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract + [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "620")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "620")
    ind.set(qn("w:hanging"), "300")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    existing_nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums + [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph(style="Legal List")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p._p.get_or_add_pPr().append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))
    return p


def add_party_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 7560])
    rows = [
        ("甲方（委托方）", "foreveryoung2023car旅乘-日本自由行包車"),
        ("法定登记名称", "________________________________（签署前填写）"),
        ("地址", "________________________________（签署前填写）"),
        ("代表人", "________________________________（签署前填写）"),
        ("联系邮箱", "foreveryoung2023car@gmail.com"),
        ("乙方（受托方）", "藍米株式会社（蓝米株式会社）"),
        ("地址", "〒106-0031 東京都港区西麻布2-13-12 早野ビル7階 ティショク西麻布 Room E"),
        ("电话", "03-6629-3857"),
        ("负责人", "川内"),
        ("联系邮箱", "sendaihazime@bluericeapp.com"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT)
        for cell_idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(cell_idx == 0))
        if idx in (0, 5):
            set_cell_shading(cells[1], PALE)
            for r in cells[1].paragraphs[0].runs:
                set_run_font(r, size=9.5, bold=True, color=NAVY)
    return table


def add_standard_table(doc, headers, rows, widths, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.text = text
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, size=font_size, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], PALE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_run_font(r, size=font_size)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.20

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(12)

    for name, size, before, after, color in (
        ("Heading 1", 14, 13, 6, NAVY),
        ("Heading 2", 11.5, 10, 4, NAVY),
        ("Heading 3", 10.5, 8, 3, ACCENT),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Legal List" not in styles:
        legal = styles.add_style("Legal List", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legal = styles["Legal List"]
    legal.font.name = FONT
    legal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    legal.font.size = Pt(10.5)
    legal.paragraph_format.space_after = Pt(4)
    legal.paragraph_format.line_spacing = 1.18


def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.34)
    section.different_first_page_header_footer = True


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("软件系统开发及交付契约书｜机密")
    set_run_font(r, size=8.2, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("foreveryoung2023car旅乘-日本自由行包車 × 藍米株式会社   ·   第 ")
    set_run_font(r, size=8, color=MUTED)
    add_page_field(p)
    r = p.add_run(" 页")
    set_run_font(r, size=8, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_header_footer(section)

    props = doc.core_properties
    props.title = "软件系统开发及交付契约书"
    props.subject = "和服预约管理系统及司机管理/派车系统"
    props.author = "foreveryoung2023car旅乘-日本自由行包車 / 藍米株式会社"
    props.keywords = "软件开发, 契约书, 和服系统, 司机管理系统"

    # Cover — proposal_centerpiece adapted for a restrained execution contract.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件系统开发及交付")
    set_run_font(r, size=27, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("契 约 书")
    set_run_font(r, size=31, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("和服预约管理系统  ·  司机管理及派车系统")
    set_run_font(r, size=13, bold=True, color=ACCENT)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2800, 6860])
    cover_rows = [
        ("合同编号", "FY-BR-________________"),
        ("合同总价", "JPY 13,464,000（含消费税）"),
        ("甲方", "foreveryoung2023car旅乘-日本自由行包車"),
        ("乙方", "藍米株式会社"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = value
        set_cell_shading(meta.cell(i, 0), LIGHT)
        for j, cell in enumerate(meta.rows[i].cells):
            for p0 in cell.paragraphs:
                p0.paragraph_format.space_after = Pt(0)
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for rr in p0.runs:
                    set_run_font(rr, size=10.5, bold=(j == 0 or i == 1))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("签署日期：________年____月____日")
    set_run_font(r, size=11, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用法律：日本法　｜　管辖：东京地方裁判所")
    set_run_font(r, size=9, color=MUTED)

    add_page_break(doc)
    add_heading(doc, "签约主体", 1)
    add_party_table(doc)
    add_body(doc, "甲乙双方基于平等、自愿及诚实信用原则，就和服预约管理系统与司机管理及派车系统的开发、整合、部署、交付及相关服务事宜，达成本契约。双方确认：附件一至附件三均为本契约不可分割的组成部分，与正文具有同等效力。")

    articles = [
        ("第一条　契约目的与项目性质", [
            "乙方根据本契约及附件约定，为甲方提供两套业务系统的需求整理、软件开发、既有功能整合、测试、部署、数据结构配置、操作资料编制及交付服务。",
            "本项目采用分阶段成果交付方式。除附件明确列为成果保证的事项外，乙方不对甲方的营业收入、客户数量、平台审核结果、第三方服务持续可用性或特定商业效果作出保证。",
            "本契约未明示包含的新增功能、重大流程变更、第三方系统深度集成、历史数据清洗迁移或持续运营服务，均依第八条办理变更。",
        ]),
        ("第二条　项目范围", [
            "项目A：和服预约管理系统。范围包括面向客户的预约及查询、退款申请、门店现场业务、客服/管理后台、订单与款项管理、权限、通知邮件、审计及Firebase相关后端功能，具体以附件一为准。",
            "项目B：司机管理及派车系统。范围包括订单派车、司机与车辆档案、地区与权限、排班、多日行程、行程状态、费用、奖惩、评价、历史记录及相关数据结构，具体以附件一为准。",
            "网站文案、照片、商标、价格规则、银行资料、业务政策及客户合法授权等业务资料由甲方提供并确认；乙方负责按确认内容实施。",
        ]),
        ("第三条　契约期限与里程碑", [
            "本契约自双方签署或盖章之日起生效。项目开始日为________年____月____日，计划最终交付日为________年____月____日。",
            "里程碑、交付物与付款节点见附件二。因甲方迟延提供资料、确认、账号、测试反馈或因第三方平台审核/故障造成的期间，应相应顺延，不视为乙方迟延。",
            "任何计划日期的调整应以电子邮件、双方认可的即时通讯记录或书面变更单确认。",
        ]),
        ("第四条　合同价款与支付", [
            "本契约总价为13,464,000日元（大写：壹仟叁佰肆拾陆万肆仟日元整），含日本消费税；适用税额及税务信息以乙方依法开具的请求书或适格发票记载为准。",
            "第一期款40%，即日元5,385,600元，于本契约签署及项目启动请求书出具后7个营业日内支付。",
            "第二期款40%，即日元5,385,600元，于两套系统主要功能提交测试环境并书面通知后7个营业日内支付。",
            "尾款20%，即日元2,692,800元，于最终验收合格或依第七条视为验收合格后7个营业日内支付。",
            "甲方以银行转账方式支付至乙方指定账户，汇款手续费由甲方承担。甲方逾期付款超过10个营业日且经催告后仍未支付的，乙方可暂停未完成服务，工期相应顺延。",
        ]),
        ("第五条　甲方义务", [
            "按时指定有决策权限的项目负责人，并及时提供准确、完整且可合法使用的业务规则、内容、素材、账号及测试数据。",
            "在约定期限内完成需求确认、测试及验收反馈。多名甲方人员意见不一致时，以甲方指定负责人最终书面意见为准。",
            "负责第三方服务账户的申请、实名认证、费用支付及持续维护，包括但不限于域名、GitHub、Firebase、Google、Gmail、图片存储及其他平台。",
            "不得要求乙方实施违法、侵权、绕过平台安全措施或违反第三方服务条款的功能。",
        ]),
        ("第六条　乙方义务", [
            "按照本契约、附件及双方确认的规格，以符合通常专业水准的方式完成开发、测试与交付，并合理报告项目进度及影响工期的事项。",
            "对开发过程中接触的甲方资料、账号及个人信息采取合理的访问控制和保密措施；仅为履行本项目目的使用。",
            "交付时提供附件二所列源代码、部署信息、数据结构说明及必要操作资料；涉及安全的密钥和密码采用安全方式另行移交。",
            "对验收期或免费缺陷修复期内确认的可复现缺陷，在合理期间内修复。",
        ]),
        ("第七条　交付、验收与缺陷处理", [
            "乙方通过测试网址、生产网址、代码仓库、压缩包或双方认可的其他方式提交交付物，并发出书面验收通知。",
            "甲方应在收到每一里程碑交付物后10个营业日内，依据附件三完成测试并一次性提交可复现的不符合项清单，载明操作步骤、预期结果、实际结果及必要截图。",
            "如交付物实质符合已确认规格，轻微视觉差异、非阻断性问题或第三方平台限制不影响验收。乙方修复阻断性或重大缺陷后，甲方应在5个营业日内复验相关项目。",
            "甲方在验收期内未提出具体书面异议，或已将交付物投入正式营业使用的，视为该里程碑验收合格；但不影响免费缺陷修复期内发现的潜在程序缺陷。",
            "需求变更、甲方数据或操作错误、未经乙方同意的代码修改、第三方服务变更/中断、设备或网络环境问题，不属于乙方免费缺陷修复范围。",
        ]),
        ("第八条　需求变更", [
            "任何超出附件一或已确认规格的新增、删除、重做、技术架构调整、批量数据处理或外部接口变更，均应提交变更申请。",
            "乙方评估变更对费用、工期、测试及既有功能的影响后，向甲方提交书面报价或变更单；双方确认后实施。未确认前，乙方可继续原范围工作。",
            "为修复安全漏洞、遵守法律或应对第三方平台强制变更所必需的调整，双方应优先协商；超出原范围的合理成本另行承担。",
        ]),
        ("第九条　知识产权与开源软件", [
            "甲方付清本契约全部价款后，乙方为本项目专门新开发并实际交付的程序代码、页面设计、文档及配置成果之可转让著作财产权转让给甲方，包括日本著作权法第27条及第28条所规定的权利。",
            "乙方在本项目开始前已拥有的通用工具、框架、模板、算法、开发方法、组件及技术诀窍，以及第三方/开源软件，不因本契约转让。对嵌入交付物且为运行所必要的乙方既有部分，乙方授予甲方永久、非独占、全球范围、免许可费的使用及为维护目的修改之权利。",
            "开源软件及第三方组件分别受其许可证约束。乙方应在合理范围内保留相关版权及许可证声明，甲方不得要求删除依法或依许可证必须保留的声明。",
            "在法律允许范围内，乙方及实际创作者对已转让成果不向甲方或甲方合法承继人行使著作者人格权。",
            "甲方提供的商标、图片、文字、业务资料、客户数据及既有成果之权利仍归甲方或原权利人所有。甲方保证其提供及指示使用的资料具有合法权源。",
        ]),
        ("第十条　账户、数据与个人信息", [
            "业务数据归甲方所有。乙方仅在履约、测试、故障排查及经甲方授权的维护范围内处理业务数据。",
            "生产账号原则上应以甲方可控制的主体申请。乙方代为设置时，应在最终交付前完成管理员权限及凭证移交。",
            "双方应遵守适用的个人信息保护法律，不得将客户、员工、司机或门店数据用于本项目以外的目的。发生可能影响个人信息或生产系统的安全事件时，知悉方应无不当延迟地通知另一方并合作处置。",
            "甲方负责确定合法的隐私政策、保存期限、数据主体告知与同意内容；乙方按确认要求实施技术配置。",
        ]),
        ("第十一条　保密义务", [
            "一方从另一方取得的、标明为保密或依其性质应合理理解为保密的技术、经营、客户、财务、账号及合同信息，均为保密信息。",
            "接收方仅为履行本契约使用保密信息，并仅向确有必要且承担不低于本条保密义务的员工、专业顾问或经许可的再委托方披露。",
            "已公开信息、接收方无保密义务而合法知悉的信息、独立开发的信息或依法必须披露的信息不在此限；依法披露时应在法律允许范围内事先通知。",
            "本条义务在本契约终止后持续5年；个人信息、账号凭证及依法应持续保护的商业秘密不受该期限限制。",
        ]),
        ("第十二条　第三方服务与再委托", [
            "系统可能依赖Firebase、Google Cloud/Sheets/GAS、Gmail、GitHub Pages、域名、图片存储、浏览器或其他第三方服务。其价格、政策、接口、审核及可用性由第三方决定。",
            "第三方服务费、域名费、短信/邮件超量费、云资源费及应用商店等费用，除附件明确包含外，由甲方另行承担。第三方停止服务或重大变更时，乙方可提出替代方案及变更报价。",
            "乙方可将部分专业工作再委托给适格人员，但仍对其履约承担责任，并应确保相关人员承担保密及个人信息保护义务。涉及大规模生产数据访问的再委托应事先告知甲方。",
        ]),
        ("第十三条　免费缺陷修复与后续维护", [
            "最终验收合格之日起90日为免费缺陷修复期。免费范围限于交付物不符合已确认规格且可复现的程序缺陷。",
            "功能新增、业务规则变化、内容更新、数据录入/清洗、第三方平台适配、甲方或第三人修改造成的问题、现场培训及持续运维不属于免费范围。",
            "免费期届满后的维护、监控、备份、值班响应、功能迭代或安全升级，由双方另行签订维护契约或按乙方报价执行。",
        ]),
        ("第十四条　陈述保证与责任限制", [
            "双方保证有权签署并履行本契约。乙方保证其明知范围内，专门开发成果不会故意侵犯第三方知识产权；甲方保证其提供的资料、数据及业务指示合法。",
            "除故意或重大过失、人身损害、违反保密/个人信息义务或知识产权侵权责任外，任一方基于本契约承担的损害赔偿总额，以甲方就引起责任的服务已实际支付且不超过本契约总价的金额为上限。",
            "在适用法律允许范围内，任何一方均不对间接损失、特别损失、逸失利益、商誉损失或可由合理备份避免的数据损失负责。",
        ]),
        ("第十五条　暂停、解除与终止后处理", [
            "一方严重违反本契约，经另一方书面催告并给予10个营业日合理补救期后仍未补救的，守约方可解除全部或部分契约。破产、停止支付、吊销营业资格或明显无法履约的，另一方可立即解除。",
            "甲方无归责于乙方之事由而提前终止时，应支付截至终止日已完成工作对应价款、不可取消的第三方成本及合理收尾费用；乙方应交付已付款对应的可用成果。",
            "契约终止后，双方应在合理期限内返还或删除对方保密资料；依法必须保存或备份系统自动保留的除外，并继续承担保密义务。",
        ]),
        ("第十六条　反社会势力排除", [
            "双方声明并保证自身及其管理人员不属于暴力团、暴力团员、关联企业、总会屋或其他反社会势力，且不利用反社会势力、不提供利益、不实施暴力要求、威胁或妨害信用/业务的行为。",
            "违反前款时，另一方可不经催告立即解除本契约，且无需赔偿因此给违约方造成的损失。",
        ]),
        ("第十七条　不可抗力", [
            "因地震、台风、火灾、战争、传染病、政府行为、大规模通信/云服务故障、罢工或其他超出合理控制范围的事件导致不能或迟延履约的，受影响方在合理通知并尽力减轻影响的范围内不承担违约责任，履行期限相应顺延。",
        ]),
        ("第十八条　通知、电子记录与签署", [
            "与本契约有关的确认、通知及变更，可通过纸面文件、电子邮件、双方认可的电子签署服务或可保存记录的业务通讯工具作出。涉及价款、范围、知识产权或解除的重大变更，应由双方有权限人员明确确认。",
            "本契约可由双方分别签署同内容文本，并可采用电子签名或盖章扫描件；各文本合并视为同一契约。",
        ]),
        ("第十九条　适用法律与争议解决", [
            "本契约的成立、效力、解释及履行适用日本法律。",
            "因本契约产生或与本契约有关的争议，双方应先本着诚信协商解决；协商不成时，以东京地方裁判所为第一审专属合意管辖法院。",
        ]),
        ("第二十条　其他", [
            "本契约及附件构成双方关于本项目的完整合意，并取代此前就同一事项作出的口头或书面沟通；但双方明确确认继续有效的保密协议除外。",
            "本契约任何条款被认定无效或不可执行，不影响其他条款效力；双方应以最接近原商业目的的有效条款替代。",
            "未尽事宜由双方依诚信原则协商，并以书面补充协议确定。补充协议与本契约不一致时，以后签署的补充协议为准。",
        ]),
    ]

    for heading, clauses in articles:
        add_heading(doc, heading, 1)
        num_id = add_numbering_definition(doc)
        for clause in clauses:
            add_numbered_paragraph(doc, clause, num_id)

    add_heading(doc, "签署页", 1)
    add_body(doc, "本契约一式两份，甲乙双方各执一份；采用电子签署时，各方持有的电子副本具有同等效力。双方确认已充分阅读并理解全部条款及附件。", align=WD_ALIGN_PARAGRAPH.LEFT, after=14)

    sig = doc.add_table(rows=6, cols=2)
    sig.style = "Table Grid"
    set_table_geometry(sig, [4830, 4830], indent=0)
    left = [
        "甲方：foreveryoung2023car旅乘-日本自由行包車",
        "法定登记名称：____________________________",
        "地址：____________________________________",
        "代表人：__________________________________",
        "签名/盖章：",
        "日期：________年____月____日",
    ]
    right = [
        "乙方：藍米株式会社",
        "地址：東京都港区西麻布2-13-12 早野ビル7階\nティショク西麻布 Room E",
        "负责人：川内",
        "电话：03-6629-3857",
        "签名/盖章：",
        "日期：________年____月____日",
    ]
    for i in range(6):
        sig.cell(i, 0).text = left[i]
        sig.cell(i, 1).text = right[i]
        if i == 0:
            set_cell_shading(sig.cell(i, 0), LIGHT)
            set_cell_shading(sig.cell(i, 1), LIGHT)
        for cell in sig.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.10
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(i == 0))
        if i == 4:
            for cell in sig.rows[i].cells:
                cell.paragraphs[0].paragraph_format.space_after = Pt(42)

    add_page_break(doc)
    add_heading(doc, "附件一　项目范围及功能清单", 1)
    add_body(doc, "本附件依据当前项目代码及业务说明整理。列明功能为本契约基准范围；实际操作名称可因界面语言或版本迭代存在不影响实质功能的差异。", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "A. 和服预约管理系统", 2)
    kimono_rows = [
        ("A1", "客户预约网站", "和服体验介绍、门店与方案展示、日期/时段/人数/方案/妆发/摄影等需求填写、折扣码验证、订金与汇款凭证流程、预约编号生成及防重复提交。"),
        ("A2", "客户自助服务", "以订单编号及联系方式查询状态；取消/退款申请、退款资料收集；适用时支持客户自助报到及状态展示。"),
        ("A3", "门店现场端", "门店登录、今日/明日订单、搜索、到店报到、人数及费用调整、现场收款/尾款、门店成本与平台结算信息。"),
        ("A4", "客服与订单后台", "登录与角色权限、订单列表/搜索/筛选/日历、订单详情编辑、款项与凭证、备注、状态流转、退款处理及客户沟通辅助。"),
        ("A5", "财务与对账", "订金、尾款、折扣、退款、超时/污损费用、平台费、门店利润/应收应付计算；对账预览、异常提示及导出。"),
        ("A6", "组织与权限", "员工账号、角色及门店范围控制，包含Owner/Admin/客服/店长/店员/会计/只读等业务角色的配置与后端校验。"),
        ("A7", "门店及业务配置", "多门店资料、付款设置、和服/妆发/摄影等项目价格与业务参数、折扣码及相关设置。"),
        ("A8", "通知与审计", "确认、提醒、退款等邮件模板与发送；关键操作审计、邮件记录及必要的后台查询。"),
        ("A9", "后端与数据", "Firebase/Firestore/Cloud Functions相关订单、用户、退款、门店、付款设置、凭证、报到、邮件及审计服务；安全规则与索引配置。"),
        ("A10", "部署与兼容", "面向现代桌面及移动浏览器的响应式页面；静态网站与Firebase相关部署配置。第三方平台配额、审核与收费不包含在固定价内。"),
    ]
    add_standard_table(doc, ["编号", "模块", "基准功能"], kimono_rows, [900, 1850, 6910], font_size=8.4)

    add_heading(doc, "B. 司机管理及派车系统", 2)
    driver_rows = [
        ("B1", "派车工作台", "订单列表、详情、搜索与筛选、地区视图、待派/已派/确认/行程中/完成/对账/取消等状态管理。"),
        ("B2", "订单与行程", "服务日期时间、起终点、行程简述、订单来源、乘客构成、语言需求、车数、车型、备注及多日行程拆分/预览。"),
        ("B3", "司机管理", "司机档案、照片、联系方式、业务标签、语言、地区/负责人、启停/休假、默认费用、年资、趟次、评分及应领统计。"),
        ("B4", "车辆管理", "车辆编号、车牌、车款、颜色、座位数、车辆分类、状态及累计使用信息。"),
        ("B5", "排班与冲突", "按日期/司机查看排班、多日行程占用、休假状态、派车冲突与可用性提示。"),
        ("B6", "费用与结算", "单次司机费用、默认日费/单程费、司机应领统计、订单奖惩、奖金/扣费及对账状态。"),
        ("B7", "评价与历史", "评分与评语、订单异动历史、状态历史、处理记录及必要的操作追踪。"),
        ("B8", "权限与地区", "派车负责人识别、所属地区访问限制、账号与角色相关界面及数据范围控制。"),
        ("B9", "数据连接", "与约定数据源/API/表格的数据读取、写入及字段映射；历史数据质量问题及大规模清洗按变更处理。"),
        ("B10", "响应式与部署", "面向调度人员的桌面工作台及必要移动端适配、部署配置与版本交付。"),
    ]
    add_standard_table(doc, ["编号", "模块", "基准功能"], driver_rows, [900, 1850, 6910], font_size=8.4)

    add_heading(doc, "C. 明确不包含事项", 2)
    excluded = [
        "第三方云服务、域名、邮件、短信、图片存储、应用商店或支付机构费用。",
        "未列明的原生iOS/Android应用、硬件采购、车载设备或GPS硬件集成。",
        "大规模历史数据录入、纠错、去重、翻译、图片拍摄、商业文案及法律政策起草。",
        "7×24小时运维、现场常驻、持续内容运营、营销投放或第三方平台审核保证。",
        "甲方提出的新增模块或因外部平台重大变更所需的重构。",
    ]
    ex_num = add_numbering_definition(doc)
    for item in excluded:
        add_numbered_paragraph(doc, item, ex_num)

    add_heading(doc, "附件二　里程碑、交付物与付款", 1)
    milestone_rows = [
        ("M1", "启动与基准确认", "生效后；计划日期：________", "需求/范围基准、账号与环境清单、实施计划", "40%\nJPY 5,385,600"),
        ("M2", "主要功能测试交付", "计划日期：________", "两套系统测试环境、主要功能、阶段测试说明", "40%\nJPY 5,385,600"),
        ("M3", "生产交付与验收", "计划日期：________", "生产版本、源代码、部署资料、数据结构/操作资料、凭证移交", "20%\nJPY 2,692,800"),
    ]
    add_standard_table(doc, ["阶段", "里程碑", "目标日期", "主要交付物", "付款"], milestone_rows, [700, 1650, 1700, 3910, 1700], font_size=8.3)
    add_body(doc, "交付介质：代码仓库访问权、压缩包、部署网址、电子文档或双方认可的其他电子方式。生产密钥、密码、恢复代码等敏感信息应通过安全渠道另行移交，不直接写入本契约。", bold_lead="交付介质：", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "附件三　验收标准与检查表", 1)
    acceptance_rows = [
        ("1", "可访问性", "约定环境可访问，关键页面无阻断性加载错误。", "□通过　□不通过"),
        ("2", "核心流程", "和服预约/查询/后台处理及司机派车/状态/排班等核心流程可按确认规格完成。", "□通过　□不通过"),
        ("3", "权限", "约定角色可访问其授权功能，关键写操作具备后端或数据源层面的权限控制。", "□通过　□不通过"),
        ("4", "数据", "新增、读取、修改及状态流转可正常保存；字段映射与金额计算符合确认规则。", "□通过　□不通过"),
        ("5", "通知", "约定邮件/通知可在第三方服务正常且配置正确时触发。", "□通过　□不通过"),
        ("6", "兼容性", "在双方约定的现代浏览器及典型桌面/手机尺寸下可正常使用。", "□通过　□不通过"),
        ("7", "交付资料", "源代码、部署说明、必要账号权限及操作资料已移交。", "□通过　□不通过"),
        ("8", "已知限制", "不影响核心营业的已知限制已记录，并明确是否纳入后续修复。", "□通过　□不通过"),
    ]
    add_standard_table(doc, ["序号", "验收项", "标准", "结果"], acceptance_rows, [700, 1500, 5860, 1600], font_size=8.4)
    add_body(doc, "验收结论：□ 合格　　□ 附条件合格（见问题清单）　　□ 不合格（存在阻断性/重大缺陷）", bold_lead="验收结论：", align=WD_ALIGN_PARAGRAPH.LEFT, after=10)
    add_body(doc, "甲方验收负责人：________________　日期：________年____月____日", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(doc, "乙方项目负责人：________________　日期：________年____月____日", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(doc, "问题清单编号/链接：____________________________________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Document-wide widow/orphan and paragraph formatting audit-friendly defaults.
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
